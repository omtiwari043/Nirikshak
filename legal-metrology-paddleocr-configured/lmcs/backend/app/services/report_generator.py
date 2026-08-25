"""
Generates downloadable compliance reports in PDF (fixed, evidentiary format)
and DOCX (editable, for officers to annotate before finalizing).
"""
from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
)

from app.config import settings

SEVERITY_COLORS = {
    "critical": colors.HexColor("#B91C1C"),
    "major": colors.HexColor("#C2410C"),
    "minor": colors.HexColor("#A16207"),
}
STATUS_COLORS = {
    "compliant": colors.HexColor("#15803D"),
    "minor_issues": colors.HexColor("#C2410C"),
    "non_compliant": colors.HexColor("#B91C1C"),
}


def _report_paths(report_id: str) -> tuple[str, str]:
    base = Path(settings.REPORT_DIR)
    return str(base / f"{report_id}.pdf"), str(base / f"{report_id}.docx")


def generate_pdf_report(
    report_id: str,
    product: dict,
    scan: dict,
    compliance: dict,
    officer: dict,
    image_path: str | None = None,
) -> str:
    pdf_path, _ = _report_paths(report_id)
    doc = SimpleDocTemplate(pdf_path, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleX", parent=styles["Title"], fontSize=16, spaceAfter=4)
    subtitle_style = ParagraphStyle("Subtitle", parent=styles["Normal"], fontSize=9, textColor=colors.grey)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=12, spaceBefore=12, spaceAfter=6)
    body = styles["Normal"]

    story = []
    story.append(Paragraph("Legal Metrology Compliance Report", title_style))
    story.append(Paragraph(
        "Generated under the Legal Metrology (Packaged Commodities) Rules, 2011 — automated pre-screening. "
        "This report is a decision-support tool; final compliance determination rests with the inspecting officer.",
        subtitle_style,
    ))
    story.append(Spacer(1, 10))

    status = compliance["status"]
    status_label = status.replace("_", " ").upper()
    status_table = Table(
        [["Report ID", report_id], ["Generated", datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")],
         ["Overall Status", status_label], ["Compliance Score", f"{compliance['score']} / 100"]],
        colWidths=[110, 350],
    )
    status_table.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("TEXTCOLOR", (1, 2), (1, 2), STATUS_COLORS.get(status, colors.black)),
        ("FONTNAME", (1, 2), (1, 2), "Helvetica-Bold"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#DDDDDD")),
    ]))
    story.append(status_table)
    story.append(Spacer(1, 10))

    story.append(Paragraph("Product & Scan Details", h2))
    product_rows = [
        ["Product Name", product.get("name", "-")],
        ["Brand", product.get("brand") or "-"],
        ["Category", product.get("category", "-")],
        ["Manufacturer (declared)", product.get("manufacturer_name") or "Not detected"],
        ["Imported", "Yes" if product.get("is_imported") else "No"],
        ["Source / Channel", product.get("source_channel") or "-"],
        ["Listing Type", scan.get("listing_type", "-")],
        ["Inspection Location", scan.get("inspection_location_text") or "-"],
        ["Scanned By", officer.get("full_name", "-")],
        ["Scan Date", scan.get("created_at", "-")],
    ]
    pt = Table(product_rows, colWidths=[150, 310])
    pt.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F3F4F6")),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#DDDDDD")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(pt)
    story.append(Spacer(1, 10))

    if image_path and os.path.exists(image_path):
        story.append(Paragraph("Photographic Evidence", h2))
        try:
            story.append(RLImage(image_path, width=170 * mm, height=110 * mm, kind="proportional"))
        except Exception:
            story.append(Paragraph("(Image could not be embedded — see attached original file.)", body))
        story.append(Spacer(1, 10))

    story.append(Paragraph(f"Violations & Non-Compliance Findings ({len(compliance['violations'])})", h2))
    if not compliance["violations"]:
        story.append(Paragraph("No violations detected by automated screening.", body))
    else:
        rows = [["Declaration", "Rule Ref.", "Type", "Severity", "Details"]]
        for v in compliance["violations"]:
            rows.append([
                Paragraph(v["declaration_title"], body),
                v.get("rule_reference") or "-",
                v["violation_type"].replace("_", " "),
                v["severity"].upper(),
                Paragraph(v.get("description") or "-", body),
            ])
        vt = Table(rows, colWidths=[95, 60, 65, 50, 190], repeatRows=1)
        style_cmds = [
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F2937")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#DDDDDD")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ]
        for i, v in enumerate(compliance["violations"], start=1):
            style_cmds.append(("TEXTCOLOR", (3, i), (3, i), SEVERITY_COLORS.get(v["severity"], colors.black)))
        vt.setStyle(TableStyle(style_cmds))
        story.append(vt)

    story.append(Spacer(1, 14))
    story.append(Paragraph(
        "Disclaimer: This report is produced by an automated image/text analysis system and is intended to "
        "assist, not replace, human judgement by a Legal Metrology enforcement officer. All findings should be "
        "physically verified before any regulatory or punitive action is taken.",
        ParagraphStyle("Disclaimer", parent=styles["Normal"], fontSize=7.5, textColor=colors.grey),
    ))

    doc.build(story)
    return pdf_path


def generate_docx_report(
    report_id: str,
    product: dict,
    scan: dict,
    compliance: dict,
    officer: dict,
) -> str:
    _, docx_path = _report_paths(report_id)
    document = Document()

    title = document.add_heading("Legal Metrology Compliance Report", level=1)
    sub = document.add_paragraph(
        "Editable draft — auto-generated by the Compliance Checking System. "
        "Officers may annotate this document before final submission."
    )
    sub.runs[0].italic = True
    sub.runs[0].font.size = Pt(9)

    document.add_paragraph(f"Report ID: {report_id}")
    document.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    status_p = document.add_paragraph()
    run = status_p.add_run(f"Overall Status: {compliance['status'].replace('_', ' ').upper()}")
    run.bold = True
    document.add_paragraph(f"Compliance Score: {compliance['score']} / 100")

    document.add_heading("Product & Scan Details", level=2)
    table = document.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    details = [
        ("Product Name", product.get("name", "-")),
        ("Brand", product.get("brand") or "-"),
        ("Category", str(product.get("category", "-"))),
        ("Manufacturer (declared)", product.get("manufacturer_name") or "Not detected"),
        ("Imported", "Yes" if product.get("is_imported") else "No"),
        ("Listing Type", scan.get("listing_type", "-")),
        ("Inspection Location", scan.get("inspection_location_text") or "-"),
        ("Scanned By", officer.get("full_name", "-")),
    ]
    for k, v in details:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = str(v)

    document.add_heading(f"Violations & Findings ({len(compliance['violations'])})", level=2)
    if not compliance["violations"]:
        document.add_paragraph("No violations detected by automated screening.")
    else:
        vtable = document.add_table(rows=1, cols=5)
        vtable.style = "Light Grid Accent 1"
        hdr = vtable.rows[0].cells
        for i, h in enumerate(["Declaration", "Rule Ref.", "Type", "Severity", "Details"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for v in compliance["violations"]:
            cells = vtable.add_row().cells
            cells[0].text = v["declaration_title"]
            cells[1].text = v.get("rule_reference") or "-"
            cells[2].text = v["violation_type"].replace("_", " ")
            cells[3].text = v["severity"].upper()
            cells[4].text = v.get("description") or "-"

    document.add_heading("Officer's Notes", level=2)
    document.add_paragraph("[Add manual observations, photographic annotations, or corrective directions here.]")

    document.add_paragraph(
        "\nDisclaimer: Automated pre-screening tool output. Verify all findings physically before "
        "regulatory action.",
    ).runs[0].font.size = Pt(8)

    document.save(docx_path)
    return docx_path
